from __future__ import annotations

from pathlib import Path

from .models import ContentUnit, SourceLocation


def import_file(path: str | Path) -> list[ContentUnit]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".xlsx":
        return import_xlsx(file_path)
    if suffix == ".docx":
        return import_docx(file_path)
    if suffix in {".md", ".txt"}:
        return import_text(file_path)

    raise ValueError(
        f"Unsupported file type: {suffix}. Supported file types: .xlsx, .docx, .md, .txt"
    )


def import_text(path: Path) -> list[ContentUnit]:
    units: list[ContentUnit] = []

    for line_number, line in enumerate(
        path.read_text(encoding="utf-8", errors="ignore").splitlines(),
        start=1,
    ):
        text = line.strip()
        if not text:
            continue

        units.append(
            ContentUnit(
                id=f"TXT-{line_number}",
                text=text,
                source=SourceLocation(
                    document=str(path),
                    document_type=path.suffix.lower().lstrip("."),
                    location={"line": line_number},
                ),
                kind="line",
            )
        )

    return units


def import_docx(path: Path) -> list[ContentUnit]:
    from docx import Document

    doc = Document(str(path))
    units: list[ContentUnit] = []

    for paragraph_number, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue

        units.append(
            ContentUnit(
                id=f"DOCX-P-{paragraph_number}",
                text=text,
                source=SourceLocation(
                    document=str(path),
                    document_type="docx",
                    location={"paragraph": paragraph_number},
                ),
                kind="paragraph",
            )
        )

    for table_number, table in enumerate(doc.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            for column_number, cell in enumerate(row.cells, start=1):
                text = cell.text.strip()
                if not text:
                    continue

                units.append(
                    ContentUnit(
                        id=f"DOCX-T-{table_number}-{row_number}-{column_number}",
                        text=text,
                        source=SourceLocation(
                            document=str(path),
                            document_type="docx",
                            location={
                                "table": table_number,
                                "row": row_number,
                                "column": column_number,
                            },
                        ),
                        kind="table_cell",
                    )
                )

    return units


def import_xlsx(path: Path) -> list[ContentUnit]:
    import openpyxl

    workbook = openpyxl.load_workbook(str(path), data_only=True)
    units: list[ContentUnit] = []

    for worksheet in workbook.worksheets:
        headers: dict[int, str] = {}

        for row_number, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            values = ["" if value is None else str(value).strip() for value in row]

            if row_number == 1:
                headers = {
                    index + 1: value
                    for index, value in enumerate(values)
                    if value
                }
                continue

            for column_number, value in enumerate(values, start=1):
                if not value:
                    continue

                units.append(
                    ContentUnit(
                        id=f"XLSX-{worksheet.title}-{row_number}-{column_number}",
                        text=value,
                        source=SourceLocation(
                            document=str(path),
                            document_type="xlsx",
                            location={
                                "worksheet": worksheet.title,
                                "row": row_number,
                                "column": column_number,
                                "header": headers.get(column_number, f"Column {column_number}"),
                            },
                        ),
                        kind="cell",
                    )
                )

    return units
